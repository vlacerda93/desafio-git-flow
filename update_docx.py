import docx
from docx.shared import Inches

doc = docx.Document('atividademicroservicos.docx')

# 1. Read Intro
with open('Introduçaomicroservicos.txt', 'r', encoding='utf-8') as f:
    intro_lines = f.read().split('\n\n')

# Find where to insert intro (before paragraph 2 "1. Engenharia...")
p2 = None
for p in doc.paragraphs:
    if "1. Engenharia de Requisitos e Organização Ágil" in p.text:
        p2 = p
        break

if p2:
    for line in intro_lines:
        text = line.strip()
        if text.startswith('## Introdução'):
            new_p = p2.insert_paragraph_before()
            run = new_p.add_run('Introdução')
            run.bold = True
            font = run.font
            font.size = docx.shared.Pt(16)
        elif text:
            p2.insert_paragraph_before(text)

# 2. Insert image (before paragraph 21 "4. Documento de...")
p21 = None
for p in doc.paragraphs:
    if "4. Documento de Fundamentação Técnica" in p.text:
        p21 = p
        break

if p21:
    # Insert an empty paragraph first
    img_p = p21.insert_paragraph_before()
    run = img_p.add_run()
    try:
        run.add_picture('diagramamicroservicosfastfood.jpeg', width=Inches(6.0))
    except Exception as e:
        print(f"Error adding jpeg: {e}")
        try:
            run.add_picture('.png', width=Inches(6.0))
        except Exception as e2:
            print(f"Error adding png: {e2}")
    
doc.save('atividademicroservicos.docx')
print("File updated successfully.")
