from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Título principal
    title = doc.add_heading('Dinâmica: O Jogo das Soluções DevSecOps', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph('Resolução dos Cenários Práticos')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_paragraph() # Espaço

    # Dados das resoluções
    resolucoes = [
        {
            "rodada": "Rodada 1: O Gargalo no Lançamento",
            "resposta": "A Ascensão do DevSecOps",
            "justificativa": "O cenário descreve o conflito clássico entre a agilidade da equipe de desenvolvimento (entregas rápidas) e a lentidão das auditorias tradicionais de segurança. A prática do DevSecOps surge exatamente para quebrar esse paradigma, automatizando processos de segurança para acompanharem a velocidade dos lançamentos ágeis sem gerar gargalos."
        },
        {
            "rodada": "Rodada 2: A Contestação da Transferência",
            "resposta": "Pilares da Segurança: CIA e Adicionais",
            "justificativa": "A situação foca diretamente em um dos pilares da segurança da informação: a Irretratabilidade (ou Não-Repúdio). O uso de logs criptografados e chaves digitais exclusivas garante a autoria de uma ação, impedindo que o cliente negue ter feito a transferência."
        },
        {
            "rodada": "Rodada 3: O Jogo de Empurra",
            "resposta": "Security by Design e Responsabilidade Compartilhada",
            "justificativa": "A fala do desenvolvedor exemplifica uma cultura defasada de silos. O princípio de Security by Design e da Responsabilidade Compartilhada estabelece que a segurança não é um passo final exclusivo da infraestrutura, mas sim um requisito embutido desde a criação (design) do código até a entrega, onde todos são responsáveis."
        },
        {
            "rodada": "Rodada 4: Revisão Manual Extrema",
            "resposta": "Segurança como Código e Testes Automatizados",
            "justificativa": "Auditar manualmente 10 mil linhas de código é um processo ineficiente, lento e sujeito a falhas humanas. A solução é implementar a Segurança como Código (como ferramentas de varredura de 'secrets' e SAST), realizando testes de forma totalmente automatizada nas esteiras de integração contínua (CI/CD)."
        },
        {
            "rodada": "Rodada 5: O Mar de Alertas",
            "resposta": "Desafios e Otimização da Gestão de Vulnerabilidades",
            "justificativa": "O excesso de falsos positivos (ruído) esconde as falhas que realmente importam. Este cenário ilustra a necessidade de otimizar a Gestão de Vulnerabilidades através da triagem de ferramentas e priorização baseada em risco crítico, garantindo que o time foque apenas nas ameaças reais."
        }
    ]

    for item in resolucoes:
        # Título da Rodada
        heading = doc.add_heading(item["rodada"], level=2)
        
        # Resposta
        p_resposta = doc.add_paragraph()
        run_label = p_resposta.add_run("Tema Relacionado: ")
        run_label.bold = True
        run_resp = p_resposta.add_run(item["resposta"])
        run_resp.font.color.rgb = RGBColor(0, 102, 204) # Azul escuro
        run_resp.bold = True
        
        # Justificativa
        p_just = doc.add_paragraph()
        p_just.add_run("Justificativa: ").bold = True
        p_just.add_run(item["justificativa"])
        
        doc.add_paragraph() # Espaço entre as rodadas

    doc.save("Respostas_Dinamica_DevSecOps.docx")
    print("Arquivo Word gerado com sucesso!")

if __name__ == "__main__":
    main()
