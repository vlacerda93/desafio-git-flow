import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Título principal
    title = doc.add_heading('Relatório: Desafio Git Flow e CI na Construção de um Dashboard O11y', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph('Participantes: Vinicius Franco e [Nome do Colega]')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_paragraph("Este documento contém as evidências de todos os 20 passos executados no desafio, seguindo as diretrizes de Git Flow, trabalho colaborativo e Integração Contínua.")

    passos = [
        {
            "titulo": "Passo 1: Estrutura Base e Push Inicial",
            "explicacao": "Inicialização do repositório local e criação da estrutura base (index.html e style.css). O primeiro push foi feito diretamente na branch 'main' para servir de fundação.",
            "comandos": "git init\ngit add .\ngit commit -m \"feat: estrutura base inicial\"\ngit branch -M main\ngit remote add origin <URL>\ngit push -u origin main",
            "prints": ["Terminal", "Pull Request (N/A para main)", "Vercel (N/A)"]
        },
        {
            "titulo": "Passo 2: Conexão com a Vercel e Criação da Develop",
            "explicacao": "Conexão do repositório à Vercel para deploy contínuo. Criação da branch 'develop' a partir da 'main' para iniciar o fluxo padrão do Git Flow.",
            "comandos": "git checkout -b develop\ngit push origin develop",
            "prints": ["Terminal", "Vercel (Status do primeiro deploy)"]
        },
        {
            "titulo": "Passo 3: Pipeline de Integração Contínua (CI)",
            "explicacao": "Criação do arquivo de CI para rodar validações no GitHub Actions em PRs direcionados para 'develop' ou 'main'.",
            "comandos": "git checkout -b feature/ci-pipeline\ngit add .github/workflows/ci.yml\ngit commit -m \"ci: adiciona pipeline de validacao\"\ngit push origin feature/ci-pipeline",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 4: Variáveis de Tema (Design System)",
            "explicacao": "Adição de variáveis de cores no CSS para criar a base do tema 'Dark Mode' (Design System) no painel.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/design-system\ngit add style.css\ngit commit -m \"feat: variaveis de tema\"\ngit push origin feature/design-system",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 5: Grid Principal de Layout",
            "explicacao": "Estruturação visual do grid que dividirá o painel em sidebar, topbar e área de conteúdo.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/main-grid\ngit add style.css\ngit commit -m \"feat: grid principal do layout\"\ngit push origin feature/main-grid",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 6: Container da Sidebar",
            "explicacao": "Criação da base e do estilo do painel lateral de navegação (Sidebar).",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/sidebar-container\ngit add index.html style.css\ngit commit -m \"feat: container da sidebar\"\ngit push origin feature/sidebar-container",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 7: Logotipo da Sidebar",
            "explicacao": "Adição do logotipo (O11y-Dash) e ícones para criar a identidade visual do dashboard.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/sidebar-logo\ngit add index.html style.css\ngit commit -m \"feat: logotipo da sidebar\"\ngit push origin feature/sidebar-logo",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 8: Menu de Navegação",
            "explicacao": "Implementação dos links do menu na sidebar (Visão Geral, Métricas, etc) e seus efeitos de hover.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/nav-menu\ngit add index.html style.css\ngit commit -m \"feat: menu de navegacao\"\ngit push origin feature/nav-menu",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 9: Barra Superior (Topbar)",
            "explicacao": "Criação da estrutura e do estilo da barra superior (Topbar) de controle.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/topbar\ngit add index.html style.css\ngit commit -m \"feat: barra superior\"\ngit push origin feature/topbar",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 10: Elementos da Topbar",
            "explicacao": "Adição de um campo de busca e avatar do perfil de usuário na Topbar.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/topbar-elements\ngit add index.html style.css\ngit commit -m \"feat: elementos da topbar\"\ngit push origin feature/topbar-elements",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 11: Container de Conteúdo Principal",
            "explicacao": "Estruturação da área que receberá os gráficos e widgets de monitoramento.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/main-content\ngit add index.html style.css\ngit commit -m \"feat: area principal de conteudo\"\ngit push origin feature/main-content",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 12: Grid de KPIs",
            "explicacao": "Criação da estrutura de grid para organizar os quatro cartões numéricos (KPIs) alinhados.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/kpi-grid\ngit add index.html style.css\ngit commit -m \"feat: grid de kpis\"\ngit push origin feature/kpi-grid",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 13: Cartão de Uptime",
            "explicacao": "Adição do primeiro widget indicando o tempo de atividade dos servidores (Uptime).",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/kpi-uptime\ngit add index.html style.css\ngit commit -m \"feat: cartao kpi uptime\"\ngit push origin feature/kpi-uptime",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 14: Cartão de Latência",
            "explicacao": "Adição do widget com informações de latência de rede (p95).",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/kpi-latency\ngit add index.html style.css\ngit commit -m \"feat: cartao kpi latencia\"\ngit push origin feature/kpi-latency",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 15: Cartão de Requisições",
            "explicacao": "Implementação do cartão que indica a quantidade de requisições por segundo.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/kpi-requests\ngit add index.html\ngit commit -m \"feat: cartao kpi requisicoes\"\ngit push origin feature/kpi-requests",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 16: Cartão de Taxa de Erro",
            "explicacao": "Inserção do cartão de alerta para acompanhamento de taxa de erros (status 5xx).",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/kpi-errors\ngit add index.html style.css\ngit commit -m \"feat: cartao kpi taxa de erros\"\ngit push origin feature/kpi-errors",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 17: Área do Gráfico Principal",
            "explicacao": "Criação de um componente visual mockado simulando um gráfico em barras de tráfego de rede vs CPU.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/main-chart\ngit add index.html style.css\ngit commit -m \"feat: mock grafico principal\"\ngit push origin feature/main-chart",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 18: Seção de Logs Recentes",
            "explicacao": "Construção da lista de eventos detalhando logs de serviços com códigos de cores dinâmicos.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b feature/recent-logs\ngit add index.html style.css\ngit commit -m \"feat: secao de logs recentes\"\ngit push origin feature/recent-logs",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 19: Preparação para Lançamento (Release)",
            "explicacao": "Criação da branch release para empacotar a versão 1.0.0, ajustes na barra de rolagem e merge final.",
            "comandos": "git checkout develop\ngit pull origin develop\ngit checkout -b release/v1.0.0\ngit add index.html style.css\ngit commit -m \"release: v1.0.0\"\ngit push origin release/v1.0.0\n# Seguido de merges no Github de release -> main e release -> develop",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        },
        {
            "titulo": "Passo 20: Correção Crítica em Produção (Hotfix)",
            "explicacao": "Criação direta da branch hotfix a partir da main para resolver a quebra de layout mobile (responsividade).",
            "comandos": "git checkout main\ngit pull origin main\ngit checkout -b hotfix/layout-mobile\ngit add style.css\ngit commit -m \"fix: responsividade no layout mobile\"\ngit push origin hotfix/layout-mobile\n# Seguido de merges de hotfix -> main e hotfix -> develop",
            "prints": ["Terminal", "Pull Request (PR) e Aprovação", "Vercel"]
        }
    ]

    for item in passos:
        doc.add_heading(item["titulo"], level=2)
        
        # Explicação
        p_exp = doc.add_paragraph()
        p_exp.add_run("1. Explicação: ").bold = True
        p_exp.add_run(item["explicacao"])
        
        # Comandos Git
        p_cmd = doc.add_paragraph()
        p_cmd.add_run("2. Comandos Git sugeridos:\n").bold = True
        cmd_run = p_cmd.add_run(item["comandos"])
        cmd_run.font.name = 'Courier New'
        
        # Placeholders de Prints
        for print_tipo in item["prints"]:
            p_print = doc.add_paragraph()
            p_print.alignment = WD_ALIGN_PARAGRAPH.CENTER
            box = p_print.add_run(f"[ COLOQUE AQUI O PRINT DE {print_tipo.upper()} ]")
            box.font.color.rgb = RGBColor(255, 0, 0)
            box.bold = True
        
        doc.add_paragraph() # Espaço

    # Adicionar explicação da Pipeline CI exigida no PDF
    doc.add_heading("4. O Papel e Funcionamento da Pipeline de CI", level=2)
    p_ci = doc.add_paragraph("Conforme exigido pelo roteiro, aqui está a explicação de cada etapa do arquivo ci.yml:\n\n"
                             "- name: CI Validate\n"
                             "Nomeia o fluxo de trabalho, facilitando a identificação visual na interface do GitHub.\n\n"
                             "- on: pull_request: branches: [develop, main]\n"
                             "Gatilhos. Dispara os testes automaticamente sempre que há um PR para develop ou main.\n\n"
                             "- runs-on: ubuntu-latest\n"
                             "Aloca uma VM Linux limpa na nuvem para executar testes sem depender do ambiente local.\n\n"
                             "- uses: actions/checkout@v3\n"
                             "Clona o código da branch do PR para dentro da VM de testes.\n\n"
                             "- run: [comandos]\n"
                             "Executa os testes ou verificações finais no código.")

    doc.save("Relatorio_Desafio_GitFlow.docx")
    print("Arquivo Relatorio_Desafio_GitFlow.docx gerado com sucesso!")

if __name__ == "__main__":
    main()
